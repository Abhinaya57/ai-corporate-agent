import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX

def read_docx_paragraphs(docx_path):
    """
    Reads all non-empty paragraphs from a .docx file.
    Returns: list of (paragraph_index, paragraph_text)
    """
    doc = Document(docx_path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append((i, text))
    return paragraphs

def add_highlighted_note(paragraph, note_text, note_id=None):
    """
    Adds an inline AI NOTE next to the paragraph, highlighted and numbered.
    """
    marker = f" [AI NOTE"
    if note_id is not None:
        marker += f" #{note_id}"
    marker += f": {note_text}]"

    run = paragraph.add_run(marker)
    try:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except Exception:
        # Fallback: do nothing if highlight not supported
        pass

def annotate_docx(docx_path, annotations, output_path):
    """
    annotations: list of (paragraph_index, comment_text)
    This implementation appends inline highlighted notes and creates an "AI NOTES" appendix.
    """
    doc = Document(docx_path)
    notes = []
    note_counter = 1

    for para_idx, comment_text in annotations:
        if para_idx < len(doc.paragraphs):
            add_highlighted_note(doc.paragraphs[para_idx], comment_text, note_id=note_counter)
            notes.append({"id": note_counter, "text": comment_text, "para_index": para_idx})
            note_counter += 1

    # Append an "AI NOTES" appendix at the end
    doc.add_page_break()
    doc.add_paragraph("AI NOTES", style='Heading 1')
    for note in notes:
        p = doc.add_paragraph()
        p.add_run(f"[AI NOTE #{note['id']}] ").bold = True
        p.add_run(f"Paragraph index: {note['para_index']}. {note['text']}")

    # Save file
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[DocUtils] Saved annotated file to: {output_path}")
